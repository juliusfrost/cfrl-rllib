import os

import docx
from docx.shared import Inches
from docx.text.run import Run

from explanations.study.builder import StudyBuilder


class DocumentStudyBuilder(StudyBuilder):
    def __init__(self, save_dir, video_dir, explanation_method, num_trials, config, name):
        super().__init__(save_dir, video_dir, explanation_method, num_trials, config, name)
        self.document = docx.Document()

    def add_explanations(self, image_path, video_path, heading_text, body_text):
        self.document.add_heading(heading_text, level=2)
        self.document.add_paragraph(body_text)
        p = self.document.add_paragraph('')
        r: Run = p.add_run()
        try:
            r.add_picture(os.path.join(video_path, image_path), height=Inches(2))
        except FileNotFoundError:
            print(f'Could not find image {image_path}')

    def add_evaluations(self, image_path, video_path, heading_text, body_text, question_text):
        self.document.add_heading(heading_text, level=2)
        self.document.add_paragraph(body_text)
        p = self.document.add_paragraph('')
        r = p.add_run()
        r.add_picture(os.path.join(video_path, image_path), height=Inches(2))
        self.document.add_paragraph(question_text)

    def save(self):
        self.document.save(os.path.join(self.save_dir, f'{self.name}.docx'))

    def trial_heading(self, trial_heading_text):
        self.document.add_heading(trial_heading_text, level=1)

    def build_intro(self, title, intro_text):
        self.document.add_heading(title, 0)
        self.document.add_paragraph(intro_text[0])
        self.document.add_paragraph(intro_text[1])
        self.document.add_paragraph(intro_text[2])
        self.document.add_paragraph()