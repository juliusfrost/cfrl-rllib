import argparse
import json
import os
import os.path

import docx
from docx.shared import Inches
from docx.text.run import Run

EXPLANATION_IDS = {
    'random': 0,
    'critical': 1,
    'counterfactual': 2,
}


def parse_args(parser_args=None):
    parser = argparse.ArgumentParser()
    parser.add_argument('--video-dir', default='videos', help='directory to load the videos from')
    parser.add_argument('--save-dir', help='directory to save output docs')
    parser.add_argument('--config', type=json.loads, default="{}", help='experiment configuration')
    parser.add_argument('--doc-id', type=str, default='000', help='id to be used as the document title')
    args = parser.parse_args(parser_args)
    # if not os.path.exists(args.video_dir):
    #     raise FileNotFoundError(f'Video directory does not exist: {args.video_dir}')
    return args


def get_explain_name(explanation_method, config):
    if explanation_method == 'random' or explanation_method == 'critical':
        def get_name(explanation_dir, trial):
            return os.path.join(explanation_dir, f'baseline_window-trial_{trial}.gif')
    elif explanation_method == 'counterfactual':
        def get_name(explanation_dir, trial):
            return os.path.join(explanation_dir,
                                f'counterfactual_window-'
                                f'{config["behavior_policy_config"]["name"]}-'
                                f't_{trial}.gif')
    else:
        raise NotImplementedError
    return get_name


def get_explain_study_text(explanation_method, config):
    text = ''
    text += 'The following video shows a segment of driver A\'s experience on highway route 1. '
    if explanation_method == 'random':
        text += 'The segments shown here are selected randomly. '
        text += 'In all parts of the video, driver A steers the car. '
    elif explanation_method == 'critical':
        text += 'The segments shown here are selected where it is critical to take certain actions. '
        text += 'In all parts of the video, driver A steers the car. '
    elif explanation_method == 'counterfactual':
        text += 'The segments shown here are selected randomly. '
        text += 'In the beginning part of the video, driver A steers the car. '
        text += 'At some point in time, driver B takes over steering ' \
                'to move the car off course than originally planned. '
        text += 'Lastly, driver A takes over control of steering. '
    else:
        raise NotImplementedError
    text += 'Your task: observe the behavior of driver A and try to detect any patterns. '
    return text


def get_eval_study_text(explanation_method, config):
    num_eval_policies = len(config['eval_config']['eval_policies'])
    text = 'The following video shows segments of driver A\'s experience on highway route 2. '
    text += 'In the beginning half of each video, driver A steers the car and shows the same thing for all videos. '
    text += 'In the ending half of each video, ' \
            'a different driver steers the car and shows different possible outcomes. '
    drivers = ['A'] + [chr(i) for i in range(ord('C'), ord('C') + num_eval_policies)]
    text += 'The drivers that steer the car in the ending half could be '
    for i, driver_name in enumerate(drivers):
        if i == len(drivers) - 1:
            text += f'or driver {driver_name}. '
        else:
            text += f'driver {driver_name}, '
    text += 'The videos are shuffled to hide which outcome was from driver A. '
    text += f'The outcomes are labeled outcome 1 through outcome {len(drivers)} from left to right. '
    text += 'Your task: remember the behavior of driver A from previous tasks,' \
            ' and guess which outcome was a result of driver A.'
    return text


def get_eval_name(eval_dir, trial):
    return os.path.join(eval_dir, f'counterfactual_window-t_{trial}.gif')


def add_explanations(document, trial, image_path, text=''):
    document.add_heading('Explanation', level=2)
    document.add_paragraph(text)
    p = document.add_paragraph('')
    r: Run = p.add_run()
    try:
        r.add_picture(image_path, height=Inches(2))
    except FileNotFoundError:
        print(f'Could not find image {image_path}')


def add_evaluations(document, trial, eval_image, text=''):
    document.add_heading('Evaluation', level=2)
    document.add_paragraph(text)
    p = document.add_paragraph('')
    r = p.add_run()
    r.add_picture(eval_image, height=Inches(2))
    document.add_paragraph('')
    document.add_paragraph('Which outcome was a result of driver A? '
                           'Write your answer in the accompanying answer sheet. ')


def build_document(save_dir, video_dir, explanation_method, num_trials, config, doc_name):
    document = docx.Document()
    document.add_heading('Explainable Reinforcement Learning User Evaluation', 0)
    document.add_paragraph(
        'The following videos show the behavior of drivers in different situations. '
        'Your task is to watch explanation videos of the driver in the yellow car '
        'and try to predict how the driver would react in different situations. '
        'We label the yellow car driver in the explanation videos as driver A. '
    )
    document.add_paragraph(
        'To measure how well you understood driver A\'s behavior, you will be given an evaluation task. '
        'In the evaluation task, you will be shown multiple videos with different drivers. '
        'In the first half of the video, there is a single driver A '
        'corresponding to the driver in the explanation videos. '
        'In the second half of the video, there is either the same driver A '
        'or a different driver that can lead to different outcomes. '
        'Your goal is to select which driver you think is the same one as the one you saw in the explanation videos '
        'by selecting the video which did not switch drivers. '
        'In other words, select the outcome obtained by driver A. '
    )
    document.add_paragraph(
        'Note that the source of the videos you see in the evaluation task '
        'may be different from that in the explanations (eg. different road conditions, different grey car behavior). '
        'We label this difference as driving on a different highway (1 or 2).'
    )
    document.add_paragraph()
    # r: Run = p.add_run()
    # r.add_text('Here are the explanations:\n')
    name_formula = get_explain_name(explanation_method, config)
    explanation_study_text = get_explain_study_text(explanation_method, config)
    eval_study_text = get_eval_study_text(explanation_method, config)
    for trial in range(num_trials):
        explanation_dir = os.path.join(video_dir, f'explain-{explanation_method}')
        eval_dir = os.path.join(video_dir, 'eval')

        document.add_heading(f'Trial {trial + 1}', level=1)
        add_explanations(document, trial, name_formula(explanation_dir, trial), explanation_study_text)
        add_evaluations(document, trial, get_eval_name(eval_dir, trial), eval_study_text)

    document.save(os.path.join(save_dir, f'{doc_name}.docx'))


def read_eval_policies(video_dir):
    eval_dir = os.path.join(video_dir, 'eval')
    eval_images = {}
    for image_file in os.listdir(eval_dir):
        if 'counterfactual_window' in image_file and '.gif' in image_file:
            base_name, ext = os.path.splitext(image_file)
            vid_type, policy, trial_name = base_name.split('-')
            trial = int(trial_name.split('_')[-1])
            if trial not in eval_images:
                eval_images[trial] = {}
            eval_images[trial][policy] = os.path.join(eval_dir, image_file)
    return eval_images


def main(parser_args=None):
    args = parse_args(parser_args)

    num_trials = args.config['eval_config']['num_trials']
    explanation_methods = args.config['explanation_method']
    assert isinstance(explanation_methods, list)

    doc_id = args.config['doc_config']['id']
    with open(os.path.join(args.save_dir, 'doc_ids.txt'), 'w') as f:
        for explanation_method in explanation_methods:
            method_doc_id = EXPLANATION_IDS[explanation_method]
            full_doc_id = f'{doc_id}-{method_doc_id:03d}'
            f.write(f'{explanation_method}, {full_doc_id},\n')
            build_document(args.save_dir, args.video_dir, explanation_method, num_trials, args.config, full_doc_id)


if __name__ == '__main__':
    main()
